import os
import csv
import base64
import html2text
import docx2txt
from docx import Document
from parsers import Parsers
from .Worker import Worker


class Processor:

    def __init__(self, worker: Worker):
        self.tmp_fldr = os.getcwd() + '/temp/'
        self.worker = worker
        self.log = self.worker.logger
        self.docx_ext = '.docx'

    @staticmethod
    def attach_extention_check(attach_name):
        ext = attach_name.split('.')[-1]
        return (
                (ext == 'docx' or ext == 'doc' or ext == 'pdf' or ext == 'xls' or ext == 'xlsx')
                and
                ('рекв' in attach_name.lower() or 'карт' in attach_name.lower())
            )

    @staticmethod
    def attach_name_check(attach_name):
        return ('тз' not in attach_name.lower().split(' ')  and 'заяв' not in attach_name.lower())

    def parse_docx(self, attachment_path):
        """Parser via python-docx"""
        full_text = ''
        document = Document(''.join(attachment_path.split('.')[:-1]) + self.docx_ext)
        if len(document.tables) > 0:
            for table_count, _ in enumerate(document.tables):
                table = document.tables[table_count]
                for i, row in enumerate(table.rows):
                    text = [cell.text + '\n' for cell in row.cells if len(cell.text) > 2]
                    full_text += '\n '.join(text)

        paragraphs = []
        for num, para in enumerate(document.paragraphs):
            paragraphs.append(para.text)
        full_text += ' '.join(paragraphs)

        if len(full_text) == 0:
            extracted_list = []
            for x in docx2txt.process(''.join(attachment_path.split('.')[:-1]) + self.docx_ext).split():
                if x not in extracted_list:
                    extracted_list.append(x)
            return ' '.join(extracted_list)

        return full_text

    @staticmethod
    def parse_csv(attachment_path):
        full_text = ''
        with open(''.join(attachment_path.split('.')[:-1]) + '.csv', newline='') as csvfile:
            spamreader = csv.reader(csvfile, delimiter=',')
            for row in spamreader:
                full_text += " " + ' '.join(row)
        return full_text

    def save_attachments_files(self, parsed_eml):
        os.makedirs(self.tmp_fldr, exist_ok=True)
        attach_names = []
        if len(parsed_eml['attachment']) > 0:
            for attach_count in range(len(parsed_eml['attachment'])):
                attach_name = parsed_eml['attachment'][attach_count]['filename']
                attach_name = attach_name.replace('.', '')[:40] + "." + attach_name.split(".")[-1]
                print(attach_name)
                if self.attach_extention_check(attach_name) and self.attach_name_check(attach_name):
                    f = open(f'{self.tmp_fldr}/{attach_name}', 'wb+')
                    f.write(base64.b64decode(parsed_eml['attachment'][attach_count]['raw']))
                    f.close()
                    attach_names.append(attach_name)
        self.log.info(f'{attach_names}')
        return attach_names

    def convert_attachment_file(self, attach_name: str):
        os.makedirs(self.tmp_fldr, exist_ok=True)
        attachment_path = f'{self.tmp_fldr}/{attach_name}'
        self.log.info(attachment_path)

        if attach_name.split('.')[-1] == 'doc':
            self.log.info('converting')
            cmd = f'lowriter --convert-to docx "{attachment_path}" --outdir "{self.tmp_fldr}"'
            os.system(cmd)
        elif attach_name.split('.')[-1] == 'pdf':
            self.log.info('converting')
            fltr = 'writer_pdf_import'
            cmd = f'libreoffice --infilter="{fltr}" --convert-to docx "{attachment_path}" --outdir "{self.tmp_fldr}"'
            os.system(cmd)
        elif attach_name.split('.')[-1] == 'xlsx' or attach_name.split('.')[-1] == 'xls':
            self.log.info('converting')
            cmd = f'unoconv -f csv "{attachment_path}"'
            os.system(cmd)
        return attachment_path

    def org_structure(self, inn, bik, r_account, corr_account, email, first, middle, last, tel, website, address):
        return {
            'inn': [inn if len(inn) == 10 or len(inn) == 12 else inn + '*'][0],
            'bik': [bik if len(bik) == 9 else bik + '*'][0],
            'r_account': [r_account if len(r_account) == 20 else r_account + '*'][0],
            'c_account': [corr_account if len(corr_account) == 20 else corr_account + '*'][0],
            'email': email,
            "first": first,
            "middle": middle,
            "last": last,
            "phone": tel,
            "website": website,
            "address": address,
        }

    def file_parsing(self, message_text, attach_texts, header_from):
        organization_dict, result = {}, []
        self.log.info(Parsers.clean_text(message_text))
        for attach_name, full_text in attach_texts.items():
            self.log.info(attach_name)
            card = Parsers.clean_text(full_text)  # full_text att_text
            self.log.info(card)

            inn = Parsers.parse_inn(card)
            bik = Parsers.parse_bik(card)
            corr_account = Parsers.parse_corr_account(card)
            r_account = Parsers.parse_r_account(card)
            email = Parsers.get_email(message_text, header_from)
            first, last, middle = Parsers.get_name(message_text, header_from)
            tel = Parsers.parse_tel(Parsers.clean_text(message_text))
            address = Parsers.get_address(full_text)
            website_check = (
                    len(Parsers.get_website(message_text)) == 0 and len(email) > 0
                    and
                    'yandex' not in email and 'gmail' not in email and 'mail.ru' not in email and 'bk.ru' not in email
            )
            website = 'www.' + email.split('@')[-1] if website_check else Parsers.get_website(message_text)

            organization_dict[str(attach_name).split('/')[-1]] = self.org_structure(
                inn, bik, r_account, corr_account, email, first, middle, last, tel, website, address
            )

            self.log.info(str(organization_dict))
            result.append(organization_dict)

            return result

    def get_message_text(self, parsed_eml):
        message_text = ''
        header_from = parsed_eml['header']['header']['from'][0]
        if len(parsed_eml['body']) > 0:
            message_text += html2text.html2text(parsed_eml['body'][-1]['content'])
        return message_text, header_from
